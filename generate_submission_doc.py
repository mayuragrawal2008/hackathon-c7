"""Build SUBMISSION.docx from SUBMISSION.md, embedding the hand-drawn diagrams."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

img_re = re.compile(r"(evidence/[A-Za-z0-9_\-]+\.(?:jpg|jpeg|png))")

for raw in Path("SUBMISSION.md").read_text().splitlines():
    line = raw.rstrip()
    if not line.strip():
        doc.add_paragraph("")
        continue
    # diagram placeholder -> embed the image
    if line.lstrip().startswith("*[") and "evidence/" in line:
        m = img_re.search(line)
        if m and Path(m.group(1)).exists():
            try:
                doc.add_picture(m.group(1), width=Inches(5.5))
            except Exception:
                doc.add_paragraph(f"[diagram: {m.group(1)}]")
        continue
    if line.startswith("# "):
        doc.add_heading(line[2:], level=0)
    elif line.startswith("## "):
        doc.add_heading(line[3:], level=1)
    elif line.startswith("### "):
        doc.add_heading(line[4:], level=2)
    elif line.strip() == "---":
        continue
    else:
        # strip basic markdown markers for clean text
        txt = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        txt = re.sub(r"\*(.+?)\*", r"\1", txt)
        txt = txt.replace("`", "")
        doc.add_paragraph(txt)

doc.save("SUBMISSION.docx")
print("SUBMISSION.docx written")
