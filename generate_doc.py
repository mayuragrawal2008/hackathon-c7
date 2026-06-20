"""Generate HANDOVER.docx — full product / handover document for Concept Check."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styling ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def h1(t):
    p = doc.add_heading(t, level=1)
    return p

def h2(t):
    return doc.add_heading(t, level=2)

def para(t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    return p

def bullet(t):
    return doc.add_paragraph(t, style="List Bullet")

def code(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return p

# ===================== TITLE =====================
title = doc.add_heading("Concept Check — Product & Handover Document", level=0)
sub = doc.add_paragraph()
r = sub.add_run("Track A · 100xEngineers Cohort 7 Code Path Hackathon")
r.italic = True
doc.add_paragraph("Owner: Mayur Agrawal   |   Live: https://hackathon-c7.onrender.com   |   "
                  "Repo: https://github.com/mayuragrawal2008/hackathon-c7")
doc.add_paragraph("Model in product: Groq · Llama 3.3 (llama-3.3-70b-versatile)")

# ===================== 1. WHAT IT IS =====================
h1("1. What the product is")
para("Concept Check tells a learner whether they truly understand a systems concept "
     "(why a backend exists, interface/API, frontend vs backend, storage/database) or "
     "only recognize the words for it. The learner explains a concept in their own "
     "words; the tool finds the single place the explanation stops being a derivation "
     "and becomes a memorized label, hands back ONE sharp follow-up question, and then "
     "checks whether the learner can now derive the answer.")
para("Core principle: understanding is judged on the causal “why / what breaks "
     "without it” — NOT on the presence of technical vocabulary. A correct "
     "plain-language answer is a pass. (See the Eddie clause in section 7.)")

# ===================== 2. ARCHITECTURE =====================
h1("2. Architecture overview")
code(
"Learner (browser: index.html + JS)\n"
"   |  login + DB reads/writes (RLS) via supabase-js\n"
"   |  /analyze and /judge calls (same origin)\n"
"   v\n"
"FastAPI backend (main.py, on Render)\n"
"   |-- Groq / Llama 3.3   (find gap, write follow-up, judge the why)\n"
"   |-- (no DB access; the LLM is boxed here)\n"
"   v\n"
"Supabase Postgres + Auth + Row-Level Security\n"
"   concepts (seed) -> sessions (user+concept) -> attempts (before/after + verdict)"
)
para("Why this split: the frontend talks to Supabase directly with the user's own "
     "token, so row-level security is enforced by the database. The backend is the "
     "ONLY place the LLM runs, which keeps the judging logic and its guardrails "
     "server-side and out of the user's reach.")

# ===================== 3. THE BOUNDARY =====================
h1("3. The boundary: deterministic vs probabilistic (Move 3)")
h2("Deterministic (fixed rules / database)")
bullet("The concept list — shipped as fixed seed data, never generated per request.")
bullet("Every session and attempt record (who, concept, both explanations, verdicts).")
bullet("The gap→result link (attempts.gap_closed) — evidence and metric at once.")
bullet("Row-level-security access rules.")
bullet("The plain pass/fail verdict stored per attempt.")
h2("Probabilistic (Groq / Llama 3.3)")
bullet("Find where the explanation stops deriving the why and becomes a label.")
bullet("Write ONE follow-up question targeting that exact gap (never gives the answer).")
bullet("Judge whether the second attempt derives the why.")
para("Who judges, and why: the LLM judges, but it is constrained to a causal test and "
     "must quote the exact proof sentence. A human-only judge is harder to fake but too "
     "slow for a live tool; the causal criterion + proof-sentence quote stop the model "
     "from rubber-stamping everyone.", )

# ===================== 4. BACKEND FUNCTIONS =====================
h1("4. Backend functions (main.py)")

h2("load_env(path='.env')")
bullet("Role: reads KEY=VALUE lines from .env into environment variables at startup.")
bullet("Used so GROQ_API_KEY is available locally; on Render the var is set in the dashboard.")

h2("groq_json(system, user) -> dict")
bullet("Role: the single choke point for every LLM call.")
bullet("Sends a system prompt (the guardrails) + the user payload to Llama 3.3.")
bullet("Forces response_format = json_object so the model must return valid JSON.")
bullet("temperature = 0.2 — low, for consistent judging.")
bullet("On any error returns HTTP 502 with the message (never silently passes).")

h2("POST /analyze  (first pass)")
bullet("Input: { concept_prompt, explanation_1 }.")
bullet("Role: decide if the learner already derived the why on the first try.")
bullet("Output: { first_pass_closed, gap_named, followup, proof_sentence }.")
bullet("If not closed, returns the named gap + one follow-up question.")

h2("POST /judge  (second pass)")
bullet("Input: { concept_prompt, explanation_1, gap_named, followup, explanation_2 }.")
bullet("Role: decide if the learner derived the why AFTER the follow-up.")
bullet("Output: { gap_closed, proof_sentence }.")
bullet("gap_closed is the load-bearing verdict stored as the gap→result link.")

h2("GET /health")
bullet("Role: liveness + sanity check. Returns model name and whether the key is present.")

h2("GET /")
bullet("Role: serves index.html (the learner-facing app).")

# ===================== 5. LLM PROMPTS & GUARDRAILS =====================
h1("5. LLM prompts and guardrails")
para("All three judging jobs share one guardrail block (RULES), then add a task-specific "
     "instruction. The guardrails are the heart of the product — they are what stop "
     "the model from being a soft judge that passes everyone.")
h2("Shared guardrails (RULES)")
bullet("Judge the causal reasoning, NOT vocabulary. Plain-language answers can pass.")
bullet("Do NOT be agreeable. Restating the term, defining it, listing features, or "
       "deflecting (“it's just how it works”) is NOT a derivation.")
bullet("Always return the exact sentence from the learner's text that proves the verdict, "
       "or an empty string if none exists.")
bullet("Return ONLY valid JSON.")
h2("ANALYZE task")
bullet("Decide first_pass_closed; if false, name the gap and write ONE follow-up that "
       "forces reasoning and never reveals the answer.")
h2("JUDGE task")
bullet("Decide gap_closed from the second explanation; quote the proof sentence or return "
       "empty if they did not derive it.")
h2("Guardrail behaviour (verified live)")
bullet("Weak/deflecting answer (“it's just a social media thing”) → gap found.")
bullet("Plain-language causal answer with zero jargon → PASS (Eddie clause).")
bullet("Deflection on the second try (“the backend just handles it”) → not "
       "closed, no proof sentence. The soft-judge trap is avoided.")

# ===================== 6. FRONTEND =====================
h1("6. Frontend functions (index.html)")
bullet("Auth handlers (signup / login / logout): Supabase email+password; the JWT is the "
       "learner's identity for row-level security.")
bullet("render(session): shows the app when logged in, the login card when not.")
bullet("loadConcepts(): reads the fixed concept list from Supabase and fills the dropdown.")
bullet("analyzeBtn handler: creates a session row, calls /analyze, stores the attempt "
       "(explanation_1, first_pass_closed, gap_named, followup). Shows the follow-up or a "
       "first-try pass.")
bullet("judgeBtn handler: calls /judge, updates the attempt with explanation_2, gap_closed, "
       "proof_sentence. Shows the verdict.")
bullet("loadHistory(): lists the learner's own past attempts — visibly demonstrates "
       "row-level security (only their rows appear).")
bullet("showVerdict / resetCards: UI helpers for the pass/fail card and proof sentence.")

# ===================== 7. DATA MODEL =====================
h1("7. Data model and row-level security (Move 4)")
h2("Tables")
bullet("concepts: id, slug, name, prompt. Fixed seed list, readable by any signed-in user.")
bullet("sessions: id, user_id (= auth.uid()), concept_id, created_at. One per learner+concept.")
bullet("attempts: id, session_id, user_id, concept_id, explanation_1, first_pass_closed, "
       "gap_named, followup, explanation_2, gap_closed, proof_sentence, created_at.")
para("The load-bearing field is attempts.gap_closed — the link from the named gap to "
     "its second-pass result. Without it you could not tell a result that worked from one "
     "that only sounded good.", bold=False)
h2("Row-level security")
bullet("Every sessions/attempts row is visible only when auth.uid() = user_id.")
bullet("Two-user test (PASSED): User A inserts an attempt; A can read it; User B reading the "
       "same table gets an empty result. Recorded in evidence/move4-rls-test.md.")

# ===================== 8. THE EDDIE CLAUSE =====================
h1("8. The Eddie clause (the guardrail that defines the product)")
para("During the by-hand Move 1 sessions, one person (“Eddie”) derived the concept "
     "correctly in plain language (“the backend is the muscle and organs; the internet is "
     "just the skeleton; the data has to live somewhere”) but lacked technical jargon. "
     "The interviewer caught himself grading on the words he expected rather than on "
     "understanding. Flagging missing vocabulary as a gap is a FALSE POSITIVE — and it is "
     "the exact soft/biased-judge failure the project warns against. The tool's judge is "
     "explicitly built to ignore jargon and test only the causal why.")

# ===================== 9. CONFIG / SECRETS =====================
h1("9. Configuration and secrets")
bullet("GROQ_API_KEY — backend env var (Render dashboard / local .env). The only "
       "server-side secret.")
bullet("SUPABASE_URL + SUPABASE_ANON_KEY — baked into index.html. The anon key is a "
       "public browser key by design; row-level security protects the data, not the key.")
bullet("SUPABASE_SERVICE_ROLE_KEY / SECRET key — admin only, never shipped to the "
       "browser, used only for setup/verification.")
bullet(".env, .venv/, .idea/ are gitignored and never committed.")
para("Security note: rotate the Groq key and the GitHub token after the event — they "
     "were shared during the build.", bold=True)

# ===================== 10. DEPLOY =====================
h1("10. Deployment runbook")
bullet("Supabase: run schema.sql in the SQL Editor (creates tables + RLS + seed). Turn OFF "
       "email confirmation for instant test logins.")
bullet("Render: connect the GitHub repo. Build: pip install -r requirements.txt. "
       "Start: uvicorn main:app --host 0.0.0.0 --port $PORT. Add env var GROQ_API_KEY.")
bullet("Free Render instances sleep when idle; first request can take ~50s to wake.")
bullet("Repo is private during the build and must be made public before the deadline "
       "(an automated routine flips it at 8 PM IST).")

# ===================== 11. LIMITATIONS =====================
h1("11. Known limitations and future work")
bullet("Accepted failure mode: a fluent but hollow causal-sounding sentence could earn a "
       "false pass. Mitigated by the required proof-sentence quote (human-inspectable).")
bullet("Concept list is fixed to systems concepts by design (keeps the deterministic anchor). "
       "Generalising to any topic would weaken the boundary.")
bullet("Single LLM provider (Groq/Llama). A second-opinion judge or a stricter rubric could "
       "further harden the verdict.")
bullet("No rate limiting or abuse controls — fine for a hackathon, needed for production.")

doc.save("HANDOVER.docx")
print("HANDOVER.docx written")
